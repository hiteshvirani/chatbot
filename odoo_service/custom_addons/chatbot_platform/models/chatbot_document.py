import os
import requests
import logging
import PyPDF2
from docx import Document
from odoo import models, fields, api
from odoo.exceptions import UserError

_logger = logging.getLogger(__name__)


class ChatbotDocument(models.Model):
    _name = 'chatbot.document'
    _description = 'Chatbot Document'
    _order = 'uploaded_at desc'

    name = fields.Char('File Name', required=True)
    chatbot_id = fields.Many2one('chatbot.chatbot', 'Chatbot', required=True, ondelete='cascade')
    
    # File information
    file_path = fields.Char('File Path')
    file_type = fields.Selection([
        ('pdf', 'PDF'),
        ('docx', 'Word Document'),
        ('txt', 'Text File'),
    ], 'File Type', required=True)
    file_size = fields.Integer('File Size (bytes)')
    content = fields.Text('Extracted Content')
    
    # Processing status
    processed = fields.Boolean('Processed', default=False)
    vector_sync_status = fields.Selection([
        ('pending', 'Pending'),
        ('synced', 'Synced'),
        ('error', 'Error')
    ], 'Sync Status', default='pending')
    
    # Timestamps
    uploaded_at = fields.Datetime('Uploaded At', default=fields.Datetime.now, readonly=True)
    updated_at = fields.Datetime('Updated At', default=fields.Datetime.now)

    @api.model
    def create(self, vals):
        # Create document record
        document = super().create(vals)
        
        # Extract text content if file_path is provided
        if document.file_path:
            content = document._extract_text()
            if content:
                document.write({'content': content})
                # Trigger sync to FastAPI
                document.sync_to_fastapi()
        
        return document

    def write(self, vals):
        # Check if file or content changed
        file_changed = 'file_path' in vals or 'content' in vals
        
        vals['updated_at'] = fields.Datetime.now()
        result = super().write(vals)
        
        if file_changed:
            # Re-extract content if file changed
            if 'file_path' in vals:
                content = self._extract_text()
                if content:
                    self.write({'content': content})
            
            # Re-sync to FastAPI
            self.sync_to_fastapi()
        
        return result

    def _extract_text(self):
        """Extract text content from uploaded file"""
        if not self.file_path or not os.path.exists(self.file_path):
            return ""
        
        try:
            if self.file_type == 'pdf':
                return self._extract_pdf_text()
            elif self.file_type == 'docx':
                return self._extract_docx_text()
            elif self.file_type == 'txt':
                return self._extract_txt_text()
            else:
                return ""
        except Exception as e:
            _logger.error(f"Error extracting text from {self.file_path}: {str(e)}")
            return ""

    def _extract_pdf_text(self):
        """Extract text from PDF file"""
        text = ""
        try:
            with open(self.file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            _logger.error(f"Error reading PDF {self.file_path}: {str(e)}")
        return text.strip()

    def _extract_docx_text(self):
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = Document(self.file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        except Exception as e:
            _logger.error(f"Error reading DOCX {self.file_path}: {str(e)}")
        return text.strip()

    def _extract_txt_text(self):
        """Extract text from TXT file"""
        text = ""
        try:
            with open(self.file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        except Exception as e:
            _logger.error(f"Error reading TXT {self.file_path}: {str(e)}")
        return text.strip()

    def sync_to_fastapi(self):
        """Send document to FastAPI for embedding"""
        if not self.content:
            _logger.warning(f"No content to sync for document {self.id}")
            return False
        
        fastapi_url = self.env['ir.config_parameter'].sudo().get_param('fastapi.url', 'http://fastapi:8000')
        internal_key = self.env['ir.config_parameter'].sudo().get_param('fastapi.internal_key')
        
        if not internal_key:
            _logger.warning("FastAPI internal key not configured")
            return False
        
        payload = {
            'chatbot_id': self.chatbot_id.id,
            'content': self.content,
            'metadata': {
                'filename': self.name,
                'file_type': self.file_type,
                'file_size': self.file_size or 0,
                'uploaded_at': self.uploaded_at.isoformat() if self.uploaded_at else None
            }
        }
        
        try:
            self.write({'vector_sync_status': 'pending'})
            
            response = requests.post(
                f"{fastapi_url}/api/internal/document/{self.id}/embed",
                json=payload,
                headers={'X-Odoo-API-Key': internal_key},
                timeout=300  # 5 minutes for large documents
            )
            
            if response.status_code == 200:
                self.write({
                    'vector_sync_status': 'synced',
                    'processed': True
                })
                _logger.info(f"Successfully synced document {self.id} to FastAPI")
                return True
            else:
                self.write({'vector_sync_status': 'error'})
                _logger.error(f"Failed to sync document {self.id}: {response.text}")
                return False
                
        except Exception as e:
            self.write({'vector_sync_status': 'error'})
            _logger.error(f"Error syncing document {self.id}: {str(e)}")
            return False

    def action_retry_sync(self):
        """Manual retry sync action"""
        self.sync_to_fastapi()
        return {
            'type': 'ir.actions.client',
            'tag': 'reload',
        }

    def unlink(self):
        """Override to delete from pgvector before Odoo deletion"""
        fastapi_url = self.env['ir.config_parameter'].sudo().get_param('fastapi.url', 'http://fastapi:8000')
        internal_key = self.env['ir.config_parameter'].sudo().get_param('fastapi.internal_key')
        
        for record in self:
            if record.processed and internal_key:
                try:
                    response = requests.delete(
                        f"{fastapi_url}/api/internal/chatbot/{record.chatbot_id.id}/source/document/{record.id}",
                        headers={'X-Odoo-API-Key': internal_key},
                        timeout=30
                    )
                    
                    if response.status_code != 200:
                        _logger.error(f"Failed to delete embeddings for document {record.id}: {response.text}")
                except Exception as e:
                    _logger.error(f"Error deleting embeddings for document {record.id}: {str(e)}")
        
        return super().unlink()
